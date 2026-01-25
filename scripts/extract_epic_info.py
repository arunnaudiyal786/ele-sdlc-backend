#!/usr/bin/env python3
"""
Epic Information Extractor

Extracts epic information from project folders and generates epic.csv file.
Based on the original extract_epic_info.py that was removed in the architecture refactor.

Usage:
    python scripts/extract_epic_info.py
"""

import csv
import logging
import re
import sys
from pathlib import Path

from docx import Document

# Add parent to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


class EpicExtractor:
    """Extract epic information from project folders and TDD documents."""

    def __init__(
        self,
        projects_dir: str = "data/raw/projects",
        output_dir: str = "data/raw/projects",
    ):
        """
        Initialize extractor.

        Args:
            projects_dir: Path to projects root directory
            output_dir: Path where epic.csv will be saved
        """
        self.projects_dir = Path(projects_dir)
        self.output_dir = Path(output_dir)
        self.epics = []

    def extract_project_id(self, folder_name: str) -> str:
        """
        Extract PRJ-XXXXX from folder name.

        Examples:
            PRJ-10051-inventory-sync-automation → PRJ-10051
            PRJ-10052-order-fulfillment → PRJ-10052

        Args:
            folder_name: Name of project folder

        Returns:
            Project ID (e.g., PRJ-10051)
        """
        match = re.match(r"(PRJ-\d+)", folder_name)
        if match:
            return match.group(1)

        logger.warning(f"Could not extract project ID from {folder_name}, using folder name")
        return folder_name

    def extract_project_title(self, folder_name: str) -> str:
        """
        Extract human-readable title from folder name.

        Examples:
            PRJ-10051-inventory-sync-automation → Inventory Sync Automation
            PRJ-10052-order-fulfillment → Order Fulfillment

        Args:
            folder_name: Name of project folder

        Returns:
            Human-readable title
        """
        # Remove PRJ-XXXXX prefix
        project_id_match = re.match(r"PRJ-\d+-(.+)", folder_name)
        if project_id_match:
            name_part = project_id_match.group(1)
            # Replace hyphens with spaces and title case
            return name_part.replace("-", " ").title()

        # Fallback: use folder name with formatting
        return folder_name.replace("-", " ").title()

    def extract_description(self, tdd_path: Path) -> str:
        """
        Extract description from TDD document.

        Strategy:
        1. Find "1.1 Purpose" heading and get next paragraph
        2. Fallback: Find first paragraph after "INTRODUCTION" or "Purpose"
        3. Last resort: Return first substantial paragraph (>50 chars)

        Args:
            tdd_path: Path to TDD.docx file

        Returns:
            Epic description text
        """
        try:
            doc = Document(str(tdd_path))

            # Strategy 1: Find "1.1 Purpose" heading
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()

                if "1.1" in text and "Purpose" in text.lower():
                    # Get next non-empty paragraph
                    for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                        next_para = doc.paragraphs[j].text.strip()
                        if len(next_para) > 20:
                            return next_para

            # Strategy 2: Find first paragraph after "INTRODUCTION" or "Purpose"
            found_intro = False
            for para in doc.paragraphs:
                text = para.text.strip()

                if "INTRODUCTION" in text.upper() or "PURPOSE" in text.upper():
                    found_intro = True
                    continue

                if found_intro and len(text) > 50:
                    return text

            # Strategy 3: First substantial paragraph
            for para in doc.paragraphs:
                text = para.text.strip()
                if len(text) > 50 and not text.isupper():  # Not a heading
                    logger.warning(f"Using fallback description extraction for {tdd_path.name}")
                    return text

            logger.error(f"Could not extract description from {tdd_path}")
            return "No description found in TDD document"

        except Exception as e:
            logger.error(f"Error reading TDD {tdd_path}: {e}")
            return f"Error extracting description: {str(e)}"

    def scan_projects(self) -> None:
        """Scan all projects in projects_dir and extract epic information."""
        if not self.projects_dir.exists():
            logger.error(f"Projects directory not found: {self.projects_dir}")
            return

        logger.info(f"Scanning projects in {self.projects_dir}")

        project_count = 0
        for project_folder in self.projects_dir.iterdir():
            # Skip files (like epic.csv itself)
            if not project_folder.is_dir():
                continue

            # Skip hidden directories
            if project_folder.name.startswith("."):
                continue

            try:
                # Extract epic information
                epic_id = self.extract_project_id(project_folder.name)
                epic_title = self.extract_project_title(project_folder.name)

                # Look for TDD document
                tdd_path = project_folder / "tdd.docx"
                if not tdd_path.exists():
                    logger.warning(f"TDD not found in {project_folder.name}, skipping")
                    continue

                epic_description = self.extract_description(tdd_path)

                # Store epic data
                self.epics.append({
                    "epicid": epic_id,
                    "epic_title": epic_title,
                    "epic_description": epic_description,
                })

                project_count += 1
                logger.info(f"✓ Extracted: {epic_id} - {epic_title}")

            except Exception as e:
                logger.error(f"Failed to process {project_folder.name}: {e}")
                continue

        logger.info(f"Successfully extracted {project_count} epics")

    def save_to_csv(self, filename: str = "epic.csv") -> None:
        """
        Save extracted epic data to CSV file.

        Args:
            filename: Name of output CSV file (saved in output_dir)
        """
        if not self.epics:
            logger.warning("No epics to save")
            return

        # Ensure output directory exists
        self.output_dir.mkdir(parents=True, exist_ok=True)

        output_path = self.output_dir / filename

        try:
            with open(output_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(
                    f,
                    fieldnames=["epicid", "epic_title", "epic_description"],
                    quoting=csv.QUOTE_MINIMAL,
                )
                writer.writeheader()
                writer.writerows(self.epics)

            logger.info(f"✓ Saved {len(self.epics)} epics to {output_path}")
            logger.info(f"  Location: {output_path.absolute()}")

        except Exception as e:
            logger.error(f"Failed to save CSV: {e}")
            raise


def main():
    """Main entry point."""
    print("=" * 80)
    print("EPIC INFORMATION EXTRACTOR")
    print("=" * 80)
    print()
    print("Extracts epic metadata from project folders and generates epic.csv")
    print()

    # Initialize extractor
    extractor = EpicExtractor(
        projects_dir="data/raw/projects",
        output_dir="data/raw/projects",
    )

    # Scan projects
    print("Scanning projects...")
    print()
    extractor.scan_projects()

    # Save to CSV
    print()
    print("Saving to CSV...")
    extractor.save_to_csv("epic.csv")

    print()
    print("=" * 80)
    print("✓ COMPLETE")
    print("=" * 80)
    print()
    print("Next steps:")
    print("  1. Review: data/raw/projects/epic.csv")
    print("  2. Run: python scripts/init_vector_db.py")
    print()


if __name__ == "__main__":
    main()
