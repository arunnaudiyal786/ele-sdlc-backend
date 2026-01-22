#!/usr/bin/env python3
"""
Generate sample documents for pipeline testing.
Creates 1 epic, 1 TDD, 1 estimation sheet, and 3 user story documents.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from datetime import datetime
import os

# Output directory
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "raw", "pipeline")

def create_epic_document():
    """Create Epic Requirements Document"""
    doc = Document()

    # Title
    title = doc.add_heading('Epic Requirements Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Epic Header
    doc.add_heading('EPIC-008: Real-Time Patient Monitoring System', level=1)

    # Metadata table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'

    metadata = [
        ('Epic ID', 'EPIC-008'),
        ('Requirement ID', 'REQ-2025-008'),
        ('JIRA ID', 'MM16790'),
        ('Priority', 'Critical'),
        ('Status', 'Planning'),
        ('Owner', 'emily.rodriguez@company.com'),
        ('Team', 'Healthcare'),
        ('Target Date', '2025-09-30'),
    ]

    for i, (key, value) in enumerate(metadata):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Description section
    doc.add_heading('Epic Description', level=2)
    description = doc.add_paragraph(
        'Build a comprehensive real-time patient monitoring system for healthcare providers '
        'to track vital signs, receive alerts for critical conditions, and coordinate care across '
        'care teams. The system must integrate with medical devices (IoT), support telemedicine '
        'workflows, ensure HIPAA compliance, and provide mobile access for clinicians.'
    )

    # Business objectives
    doc.add_heading('Business Objectives', level=2)
    objectives = [
        'Enable real-time monitoring of patient vital signs from hospital beds and remote locations',
        'Reduce response time to critical patient events by 60%',
        'Support telemedicine consultations with integrated vitals display',
        'Improve care coordination across multi-disciplinary teams',
        'Meet all HIPAA and medical device integration regulations'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # Acceptance criteria
    doc.add_heading('Acceptance Criteria', level=2)
    criteria = [
        'Integration with at least 5 major medical device manufacturers (Philips, GE, Medtronic)',
        'Real-time vitals streaming with <2 second latency',
        'Configurable alert thresholds per patient condition',
        'Mobile app for iOS and Android with offline alert queue',
        'Complete audit trail for all patient data access',
        ' 99.99% uptime SLA',
        'HIPAA compliance certification completed'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')

    # Technical requirements
    doc.add_heading('Technical Requirements', level=2)
    tech_reqs = [
        'HL7 FHIR R4 compliance for healthcare data exchange',
        'MQTT protocol for IoT device communication',
        'WebSocket support for real-time web dashboard',
        'Integration with existing EHR systems',
        'End-to-end encryption for all patient data transmission',
        'Multi-tenancy support for different healthcare facilities'
    ]
    for req in tech_reqs:
        doc.add_paragraph(req, style='List Bullet')

    # Dependencies
    doc.add_heading('Dependencies', level=2)
    doc.add_paragraph('Authentication Service (OAuth 2.0/OIDC)', style='List Bullet')
    doc.add_paragraph('Notification Service (multi-channel alerts)', style='List Bullet')
    doc.add_paragraph('Analytics Platform (historical trending)', style='List Bullet')
    doc.add_paragraph('EHR Integration Layer', style='List Bullet')

    # Save
    filename = os.path.join(OUTPUT_DIR, '1_epic_patient_monitoring.docx')
    doc.save(filename)
    print(f"✓ Created: {filename}")


def create_tdd_document():
    """Create Technical Design Document"""
    doc = Document()

    # Title
    title = doc.add_heading('Technical Design Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TDD Header
    doc.add_heading('TDD-008: Real-Time Patient Monitoring Architecture', level=1)

    # Metadata
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    metadata = [
        ('TDD ID', 'TDD-008'),
        ('Epic ID', 'EPIC-008'),
        ('Estimation ID', 'EST-008'),
        ('Version', '1.0'),
        ('Status', 'Draft'),
        ('Author', 'emily.rodriguez@company.com'),
    ]

    for i, (key, value) in enumerate(metadata):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Overview
    doc.add_heading('1. Technical Overview', level=2)
    doc.add_paragraph(
        'The Real-Time Patient Monitoring System is a distributed architecture designed to '
        'collect, process, and visualize patient vital signs from medical IoT devices. The system '
        'uses an event-driven microservices architecture with real-time streaming, complex event '
        'processing for alerts, and FHIR-compliant data storage.'
    )

    # Architecture
    doc.add_heading('2. Architecture Pattern', level=2)
    doc.add_paragraph('Event-Driven Microservices + Pub/Sub Pattern + CEP (Complex Event Processing)')
    doc.add_paragraph()

    # Technical components
    doc.add_heading('3. Technical Components', level=2)
    components = [
        ('Python', 'Backend services and CEP engine'),
        ('FastAPI', 'REST API layer'),
        ('PostgreSQL + TimescaleDB', 'Time-series vitals storage'),
        ('Apache Kafka', 'Event streaming backbone'),
        ('MQTT Broker (Mosquitto)', 'IoT device communication'),
        ('Redis', 'Real-time alert cache and session storage'),
        ('React + TypeScript', 'Web dashboard'),
        ('React Native', 'Mobile apps (iOS/Android)'),
        ('WebSockets', 'Real-time dashboard updates'),
        ('HL7 FHIR SDK', 'Healthcare data standards'),
    ]
    for component, description in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)

    # Design decisions
    doc.add_heading('4. Design Decisions', level=2)
    decisions = [
        'Use MQTT for device communication due to low bandwidth and reliability requirements',
        'TimescaleDB extension for PostgreSQL to handle time-series vitals data efficiently',
        'Apache Kafka for event streaming with patient vitals topics partitioned by facility',
        'Complex Event Processing (CEP) engine for threshold-based alerting with sliding windows',
        'WebSocket connections per care team for real-time dashboard updates',
        'FHIR Observation resources for vitals storage ensuring EHR interoperability',
        'Circuit breaker pattern for device connectivity failures',
        'Redis pub/sub for mobile push notification fan-out'
    ]
    for decision in decisions:
        doc.add_paragraph(decision, style='List Bullet')

    # Dependencies
    doc.add_heading('5. System Dependencies', level=2)
    deps = [
        'auth-service: OAuth 2.0 authentication for clinicians',
        'notification-service: Multi-channel alerts (SMS, push, email)',
        'ehr-integration: FHIR API gateway to hospital EHR systems',
        'audit-service: HIPAA-compliant audit logging',
        'analytics-platform: Historical trending and reporting'
    ]
    for dep in deps:
        doc.add_paragraph(dep, style='List Bullet')

    # Security
    doc.add_heading('6. Security Considerations', level=2)
    security = [
        'HIPAA compliance mandatory: PHI encryption at rest (AES-256) and in transit (TLS 1.3)',
        'Device authentication via X.509 certificates for MQTT connections',
        'Clinician access control via RBAC with patient assignment validation',
        'All vitals data access logged for HIPAA audit requirements',
        'Data retention policies enforced: 7 years for medical records',
        'Penetration testing required before production deployment',
        'Medical device integration must comply with FDA guidance on cybersecurity'
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')

    # Performance
    doc.add_heading('7. Performance Requirements', level=2)
    perf = [
        'Vitals ingestion: <2 second end-to-end latency from device to dashboard',
        'Alert generation: <1 second from threshold breach to notification',
        'Dashboard update rate: 1Hz (one update per second) for real-time vitals',
        'Support 10,000 concurrent patient monitoring sessions',
        'System uptime: 99.99% SLA (less than 53 minutes downtime per year)',
        'Mobile app offline mode: Queue alerts for up to 24 hours',
        'Historical query: <3 seconds for 30-day patient vitals timeline'
    ]
    for item in perf:
        doc.add_paragraph(item, style='List Bullet')

    # Data flow
    doc.add_heading('8. Data Flow', level=2)
    doc.add_paragraph(
        '1. Medical devices publish vitals to MQTT broker (topic: /facility/{id}/patient/{id}/vitals)\n'
        '2. MQTT-Kafka bridge forwards events to Kafka vitals topic\n'
        '3. Stream processor consumes events, validates, and stores in TimescaleDB\n'
        '4. CEP engine evaluates alert rules on vitals stream (sliding windows)\n'
        '5. Alerts published to alerts topic, consumed by notification service\n'
        '6. WebSocket gateway pushes vitals to connected dashboards in real-time\n'
        '7. FHIR API exposes Observation resources for EHR integration'
    )

    # Save
    filename = os.path.join(OUTPUT_DIR, '2_tdd_patient_monitoring.docx')
    doc.save(filename)
    print(f"✓ Created: {filename}")


def create_estimation_sheet():
    """Create Estimation Spreadsheet"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Effort Estimation"

    # Header row
    headers = [
        'Estimation ID', 'Epic ID', 'Module ID', 'Task Description', 'Complexity',
        'Dev Effort (hours)', 'QA Effort (hours)', 'Total Effort (hours)',
        'Story Points', 'Risk Level', 'Estimation Method', 'Confidence Level',
        'Estimated By', 'Estimation Date', 'Additional Parameters'
    ]

    # Style header
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Data row
    data = [
        'EST-008',
        'EPIC-008',
        'MOD-PTM-001',
        'Build real-time patient monitoring system with IoT device integration, MQTT message broker, '
        'complex event processing for alerts, WebSocket-based dashboard, mobile apps (iOS/Android), '
        'FHIR-compliant vitals storage, and HIPAA audit logging. Includes integration with 5 medical device manufacturers.',
        'Large',
        240,  # Dev effort
        95,   # QA effort
        335,  # Total
        42,   # Story points
        'High',
        'Planning Poker',
        'Medium',
        'senior.architect@company.com',
        '2025-02-15',
        '{"hipaa_compliance": true, "fhir_version": "R4", "mqtt_broker": "Mosquitto", '
        '"supported_devices": ["Philips", "GE", "Medtronic", "Siemens", "Mindray"], '
        '"platforms": ["web", "iOS", "Android"], "uptime_sla": "99.99%"}'
    ]

    for col, value in enumerate(data, start=1):
        cell = ws.cell(row=2, column=col, value=value)
        if col in [6, 7, 8, 9]:  # Numeric columns
            cell.alignment = Alignment(horizontal="right")
        else:
            cell.alignment = Alignment(horizontal="left", wrap_text=True)

    # Adjust column widths
    column_widths = [15, 12, 15, 60, 12, 18, 18, 20, 15, 12, 18, 18, 30, 18, 50]
    for col, width in enumerate(column_widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = width

    # Set row height for data
    ws.row_dimensions[2].height = 100

    # Save
    filename = os.path.join(OUTPUT_DIR, '3_estimation_patient_monitoring.xlsx')
    wb.save(filename)
    print(f"✓ Created: {filename}")


def create_user_story_document(story_num, story_data):
    """Create User Story Document"""
    doc = Document()

    # Title
    title = doc.add_heading(f'User Story - {story_data["title"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Story Header
    doc.add_heading(f'{story_data["id"]}: {story_data["title"]}', level=1)

    # Metadata
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'

    metadata = [
        ('Story ID', story_data['id']),
        ('Epic ID', 'EPIC-008'),
        ('Estimation ID', 'EST-008'),
        ('TDD ID', 'TDD-008'),
        ('Issue Type', 'Story'),
        ('Assignee', story_data['assignee']),
        ('Status', story_data['status']),
        ('Story Points', str(story_data['points'])),
        ('Priority', story_data['priority']),
    ]

    for i, (key, value) in enumerate(metadata):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Story description
    doc.add_heading('Description', level=2)
    doc.add_paragraph(story_data['description'])

    # Acceptance criteria
    doc.add_heading('Acceptance Criteria', level=2)
    for i, criterion in enumerate(story_data['acceptance_criteria'], start=1):
        doc.add_paragraph(f'{i}. {criterion}', style='List Number')

    # Technical notes
    if 'technical_notes' in story_data:
        doc.add_heading('Technical Notes', level=2)
        for note in story_data['technical_notes']:
            doc.add_paragraph(note, style='List Bullet')

    # Dependencies
    if 'dependencies' in story_data:
        doc.add_heading('Dependencies', level=2)
        for dep in story_data['dependencies']:
            doc.add_paragraph(dep, style='List Bullet')

    # Save
    filename = os.path.join(OUTPUT_DIR, f'{story_num}_story_{story_data["filename"]}.docx')
    doc.save(filename)
    print(f"✓ Created: {filename}")


def create_all_user_stories():
    """Create all 3 user story documents"""

    # Story 1: IoT Device Integration
    story1 = {
        'id': 'MMO-12352',
        'title': 'Build IoT Device Integration Layer',
        'filename': 'iot_integration',
        'assignee': 'dev.william@company.com',
        'status': 'To Do',
        'points': 13,
        'priority': 'Critical',
        'description': (
            'Implement MQTT-based integration layer for medical IoT devices supporting vitals '
            'streaming from Philips, GE Healthcare, Medtronic, Siemens, and Mindray patient monitors. '
            'Include device authentication via X.509 certificates, message validation, protocol translation, '
            'and fault-tolerant message buffering.'
        ),
        'acceptance_criteria': [
            'MQTT broker (Mosquitto) deployed with TLS 1.3 and certificate-based device auth',
            'Device adapters for 5 manufacturers parse proprietary protocols to standard vitals format',
            'Vitals messages validated against HL7 FHIR Observation schema',
            'Message buffering handles up to 1 hour of device disconnect',
            'Device connection status monitored with heartbeat mechanism',
            'End-to-end latency from device to Kafka topic is <2 seconds',
            'Unit tests cover all device adapters with 90% coverage'
        ],
        'technical_notes': [
            'Use paho-mqtt library for MQTT client connections',
            'Implement adapter pattern for device-specific protocol handlers',
            'Store device certificates in HashiCorp Vault',
            'Use Kafka Connect MQTT source connector for bridge to Kafka',
            'Implement circuit breaker for device connectivity failures'
        ],
        'dependencies': [
            'Network team: VPN/firewall rules for device connectivity',
            'Security team: X.509 certificate authority setup',
            'Infrastructure: MQTT broker deployment and monitoring'
        ]
    }

    # Story 2: Complex Event Processing Engine
    story2 = {
        'id': 'MMO-12353',
        'title': 'Implement Complex Event Processing for Alerts',
        'filename': 'cep_alerts',
        'assignee': 'dev.william@company.com',
        'status': 'To Do',
        'points': 13,
        'priority': 'Critical',
        'description': (
            'Build Complex Event Processing (CEP) engine for real-time patient vitals analysis and '
            'threshold-based alerting. Support sliding windows, pattern detection, and configurable '
            'alert rules per patient condition. Generate alerts for critical conditions with <1 second latency.'
        ),
        'acceptance_criteria': [
            'CEP engine consumes vitals events from Kafka with patient context',
            'Support sliding window queries (e.g., average heart rate over 5 minutes)',
            'Configurable alert rules: threshold breach, trend detection, missing data',
            'Alert severity levels: Critical, Warning, Info',
            'Alert deduplication prevents notification storms',
            'Alert generation latency <1 second from threshold breach',
            'Alert events published to Kafka alerts topic',
            'Integration tests validate complex alerting scenarios'
        ],
        'technical_notes': [
            'Consider Apache Flink for stateful stream processing',
            'Use Flink SQL for sliding window queries and pattern matching',
            'Store alert rules in PostgreSQL with versioning',
            'Implement alert suppression logic (e.g., no repeat within 5 minutes)',
            'Use Flink state backend with RocksDB for fault tolerance'
        ],
        'dependencies': [
            'IoT Integration Layer: Vitals events in Kafka',
            'Notification Service: Alert delivery channels',
            'Clinical team: Alert threshold definitions per condition'
        ]
    }

    # Story 3: Real-Time Dashboard
    story3 = {
        'id': 'MMO-12354',
        'title': 'Create Real-Time Patient Monitoring Dashboard',
        'filename': 'dashboard_ui',
        'assignee': 'dev.sophia@company.com',
        'status': 'To Do',
        'points': 13,
        'priority': 'High',
        'description': (
            'Build web-based real-time patient monitoring dashboard for care teams with live vitals display, '
            'alert notifications, patient list management, and historical trending charts. Support multi-patient '
            'monitoring with 1Hz update rate via WebSocket connections.'
        ),
        'acceptance_criteria': [
            'Patient list view with real-time vitals summary and alert indicators',
            'Detail view shows live vitals charts (heart rate, BP, SpO2, temp, resp rate)',
            'WebSocket connection maintains 1Hz update rate for selected patient',
            'Alert notifications appear as toast messages with sound',
            'Historical trends: 24-hour vitals timeline with zoom/pan',
            'Responsive design: desktop, tablet support',
            'Role-based access: only show assigned patients',
            'Dashboard loads in <2 seconds with 50 active patients'
        ],
        'technical_notes': [
            'Use React with TypeScript and Recharts for vitals visualization',
            'WebSocket library: socket.io-client for real-time updates',
            'Implement optimistic UI updates with server reconciliation',
            'Use React Query for patient data caching',
            'Vitals chart: use time-series line charts with threshold bands',
            'Implement WebSocket reconnection logic with exponential backoff'
        ],
        'dependencies': [
            'WebSocket Gateway: Real-time vitals streaming API',
            'Auth Service: RBAC patient assignment validation',
            'Design team: UI/UX mockups for dashboard layout'
        ]
    }

    create_user_story_document('4', story1)
    create_user_story_document('5', story2)
    create_user_story_document('6', story3)


def main():
    """Generate all sample documents"""
    print(f"\nGenerating sample documents in: {OUTPUT_DIR}\n")

    # Create all documents
    create_epic_document()
    create_tdd_document()
    create_estimation_sheet()
    create_all_user_stories()

    print(f"\n✅ Successfully created 6 sample documents in {OUTPUT_DIR}")
    print("\nFiles created:")
    print("  1. 1_epic_patient_monitoring.docx")
    print("  2. 2_tdd_patient_monitoring.docx")
    print("  3. 3_estimation_patient_monitoring.xlsx")
    print("  4. 4_story_iot_integration.docx")
    print("  5. 5_story_cep_alerts.docx")
    print("  6. 6_story_dashboard_ui.docx")


if __name__ == '__main__':
    main()
