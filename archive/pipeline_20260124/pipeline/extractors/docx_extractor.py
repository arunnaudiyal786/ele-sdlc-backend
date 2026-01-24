"""
DOCX document extractor using python-docx.

Extracts structured data from Word documents including:
- Document properties (title, author, dates)
- Headings with hierarchy
- Paragraphs with style information
- Tables as list of dicts
- Lists (bulleted and numbered)
- Pattern detection (Jira IDs, emails, dates, key-value pairs)
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline.extractors.base import (
    BaseExtractor,
    DocumentMetadata,
    ExtractedData,
    ExtractedField,
    ExtractedList,
    ExtractedTable,
)


class DocxExtractor(BaseExtractor):
    """
    Extracts structured data from DOCX files.

    Uses python-docx to parse Word documents and extract:
    - Document metadata (title, author, dates)
    - Headings organized by hierarchy
    - Paragraph content with style detection
    - Tables converted to list of dictionaries
    - Bulleted and numbered lists
    - Pattern detection for Jira IDs, emails, dates
    """

    # Regex patterns for common data types
    JIRA_ID_PATTERN = re.compile(r"\b([A-Z]+-\d+|MM\d+)\b")
    EMAIL_PATTERN = re.compile(r"\b[\w\.-]+@[\w\.-]+\.\w+\b")
    DATE_PATTERNS = [
        re.compile(r"\b\d{4}-\d{2}-\d{2}\b"),  # ISO: 2025-01-15
        re.compile(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b"),  # US: 1/15/2025
        re.compile(r"\b\d{1,2}-\d{1,2}-\d{2,4}\b"),  # EU: 15-01-2025
        re.compile(
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b",
            re.IGNORECASE,
        ),  # Jan 15, 2025
    ]
    KEY_VALUE_PATTERN = re.compile(r"^([A-Za-z][A-Za-z\s]{2,30}):\s*(.+)$")

    def get_supported_extensions(self) -> List[str]:
        """Return supported file extensions."""
        return [".docx"]

    async def extract(self, file_path: Path) -> ExtractedData:
        """
        Extract structured data from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractedData with all extracted content

        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If file is not a valid DOCX
        """
        self._validate_file(file_path)

        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted DOCX file: {file_path}")

        # Initialize extraction result
        extraction = ExtractedData()

        # Extract metadata
        extraction.metadata = self._extract_metadata(doc, file_path)

        # Extract content
        raw_content_parts: List[str] = []
        current_section = "Introduction"
        sections: Dict[str, List[str]] = {current_section: []}

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            raw_content_parts.append(text)

            # Check for headings
            if para.style and para.style.name.startswith("Heading"):
                current_section = text
                if current_section not in sections:
                    sections[current_section] = []
                continue

            # Add to current section
            sections[current_section].append(text)

            # Detect key-value pairs
            kv_match = self.KEY_VALUE_PATTERN.match(text)
            if kv_match:
                key = kv_match.group(1).strip()
                value = kv_match.group(2).strip()
                extraction.key_value_pairs[key] = value
                extraction.fields[key.lower().replace(" ", "_")] = ExtractedField(
                    name=key,
                    value=value,
                    confidence=0.9,
                    source_location=f"Section: {current_section}",
                )

            # Detect patterns
            extraction.jira_ids.extend(self.JIRA_ID_PATTERN.findall(text))
            extraction.emails.extend(self.EMAIL_PATTERN.findall(text))
            for pattern in self.DATE_PATTERNS:
                extraction.dates.extend(pattern.findall(text))

        # Build raw content and sections
        extraction.raw_content = "\n".join(raw_content_parts)
        extraction.raw_sections = {
            section: "\n".join(content) for section, content in sections.items()
        }

        # Extract tables
        for i, table in enumerate(doc.tables):
            extracted_table = self._extract_table(table, i)
            if extracted_table:
                extraction.tables.append(extracted_table)

        # Extract lists (check for list formatting in paragraphs)
        extraction.lists = self._extract_lists(doc)

        # Deduplicate pattern matches
        extraction.jira_ids = list(set(extraction.jira_ids))
        extraction.emails = list(set(extraction.emails))
        extraction.dates = list(set(extraction.dates))

        # Calculate overall confidence
        extraction.overall_confidence = self._calculate_confidence(extraction)

        return extraction

    def _extract_metadata(self, doc: Document, file_path: Path) -> DocumentMetadata:
        """Extract document metadata from core properties."""
        core_props = doc.core_properties

        return DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size=file_path.stat().st_size,
            file_type="docx",
            title=core_props.title if core_props.title else None,
            author=core_props.author if core_props.author else None,
            created_date=core_props.created if core_props.created else None,
            modified_date=core_props.modified if core_props.modified else None,
        )

    def _extract_table(
        self, table: Table, table_index: int
    ) -> Optional[ExtractedTable]:
        """
        Extract a table into structured format.

        Args:
            table: python-docx Table object
            table_index: Index of table in document

        Returns:
            ExtractedTable or None if table is empty
        """
        rows_data: List[Dict[str, Any]] = []
        headers: List[str] = []

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]

            # First row with content becomes headers
            if row_idx == 0:
                headers = cells if any(cells) else [f"Column_{i}" for i in range(len(cells))]
                continue

            # Create row dict using headers as keys
            if any(cells):  # Skip empty rows
                row_dict = {}
                for col_idx, cell_value in enumerate(cells):
                    if col_idx < len(headers):
                        row_dict[headers[col_idx]] = cell_value
                    else:
                        row_dict[f"Column_{col_idx}"] = cell_value
                rows_data.append(row_dict)

        if not rows_data:
            return None

        return ExtractedTable(
            headers=headers,
            rows=rows_data,
            source_location=f"Table {table_index + 1}",
            confidence=0.95,
        )

    def _extract_lists(self, doc: Document) -> List[ExtractedList]:
        """
        Extract bulleted and numbered lists from document.

        Looks for paragraphs with list formatting.
        """
        lists: List[ExtractedList] = []
        current_list: List[str] = []
        current_heading: Optional[str] = None
        in_list = False

        for para in doc.paragraphs:
            text = para.text.strip()

            # Check if paragraph has list style
            is_list_item = False
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                is_list_item = any(
                    marker in style_name
                    for marker in ["list", "bullet", "number"]
                )

            # Also check for common list markers
            if not is_list_item and text:
                is_list_item = bool(
                    re.match(r"^[\u2022\u2023\u25E6\u2043\-\*]\s", text)  # Bullet chars
                    or re.match(r"^\d+[\.\)]\s", text)  # Numbered
                    or re.match(r"^[a-z][\.\)]\s", text, re.IGNORECASE)  # Lettered
                )

            if is_list_item:
                if not in_list:
                    in_list = True
                # Clean list markers from text
                clean_text = re.sub(
                    r"^[\u2022\u2023\u25E6\u2043\-\*\d+a-zA-Z][\.\)]*\s*", "", text
                )
                if clean_text:
                    current_list.append(clean_text)
            else:
                if in_list and current_list:
                    # Save completed list
                    lists.append(
                        ExtractedList(
                            items=current_list.copy(),
                            list_type="bullet",
                            heading=current_heading,
                        )
                    )
                    current_list = []
                    in_list = False

                # Track headings for context
                if para.style and para.style.name.startswith("Heading"):
                    current_heading = text

        # Don't forget last list
        if current_list:
            lists.append(
                ExtractedList(
                    items=current_list,
                    list_type="bullet",
                    heading=current_heading,
                )
            )

        return lists

    def _calculate_confidence(self, extraction: ExtractedData) -> float:
        """
        Calculate overall extraction confidence.

        Based on:
        - Amount of content extracted
        - Number of structured elements found
        - Pattern matches
        """
        confidence = 0.5  # Base confidence

        # Content length bonus
        if len(extraction.raw_content) > 100:
            confidence += 0.1
        if len(extraction.raw_content) > 500:
            confidence += 0.1

        # Structured elements bonus
        if extraction.tables:
            confidence += 0.1
        if extraction.key_value_pairs:
            confidence += 0.1

        # Pattern matches bonus
        if extraction.jira_ids:
            confidence += 0.05
        if extraction.emails:
            confidence += 0.05

        return min(confidence, 1.0)
