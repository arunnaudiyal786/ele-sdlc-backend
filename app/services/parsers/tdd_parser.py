"""
TDD Parser

Generic Word document text extractor for LLM context.
Extracts all text content from TDD.docx files without schema assumptions.
"""

import logging
import re
from pathlib import Path
from typing import List

from docx import Document
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


# ===== Pydantic Models =====


class TableContent(BaseModel):
    """Extracted content from a single table"""

    table_number: int = Field(..., description="Table index in document (0-based)")
    context_heading: str = Field("", description="Heading before the table")
    text_content: str = Field("", description="Flat text representation of table")
    row_count: int = Field(0, description="Number of rows including header")


class TDDDocument(BaseModel):
    """Generic extracted document structure"""

    project_id: str = Field(..., description="Project ID from folder name")
    file_name: str = Field(..., description="Source file name")
    paragraph_count: int = Field(0, description="Number of paragraphs")
    table_count: int = Field(0, description="Number of tables")
    tables: List[TableContent] = Field(default_factory=list, description="Tables as flat text")
    full_text: str = Field("", description="All paragraphs combined as flat text for LLM")


# ===== Parser Class =====


class TDDParser:
    """Generic Word document parser - extracts all text without schema assumptions"""

    async def parse(self, tdd_path: Path) -> TDDDocument:
        """
        Parse any Word document into flat text format.

        Args:
            tdd_path: Path to Word document (.docx)

        Returns:
            TDDDocument with extracted text from all paragraphs and tables

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file cannot be parsed
        """
        if not tdd_path.exists():
            raise FileNotFoundError(f"File not found: {tdd_path}")

        logger.info(f"Parsing TDD document: {tdd_path.name}")

        try:
            doc = Document(str(tdd_path))
        except Exception as e:
            raise ValueError(f"Failed to open Word document: {e}")

        project_id = self._extract_project_id(tdd_path)

        # Extract paragraphs as flat text
        full_text = self._extract_full_text(doc)
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])

        # Extract tables as flat text
        tables = self._extract_tables(doc)

        logger.info(
            f"Extracted {paragraph_count} paragraphs, "
            f"{len(tables)} tables, "
            f"total {len(full_text)} characters"
        )

        return TDDDocument(
            project_id=project_id,
            file_name=tdd_path.name,
            paragraph_count=paragraph_count,
            table_count=len(tables),
            tables=tables,
            full_text=full_text,
        )

    def _extract_full_text(self, doc: Document) -> str:
        """
        Extract all text from document paragraphs.

        Preserves paragraph structure with double newlines between paragraphs.
        """
        text_parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_tables(self, doc: Document) -> List[TableContent]:
        """
        Extract all tables as flat text.

        Each row becomes a line with cell values separated by " | ".
        Tables include context from preceding headings.
        """
        tables: List[TableContent] = []
        last_heading = ""

        # Track headings before tables
        for element in doc.element.body:
            if element.tag.endswith("p"):
                # Paragraph - check if it's a heading
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para:
                    style = para.style.name if para.style else ""
                    if "Heading" in style or self._is_heading(para.text):
                        last_heading = para.text.strip()

            elif element.tag.endswith("tbl"):
                # Table
                table_obj = next((t for t in doc.tables if t._element == element), None)
                if table_obj and len(table_obj.rows) > 0:
                    table_content = self._extract_table_text(
                        table_obj, len(tables), last_heading
                    )
                    tables.append(table_content)

        return tables

    def _extract_table_text(
        self, table, table_num: int, context: str
    ) -> TableContent:
        """
        Extract a single table as flat text.

        Each row becomes a line with cell values separated by " | ".
        """
        text_lines: List[str] = []

        for row in table.rows:
            cell_values = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    cell_values.append(text)
            if cell_values:
                text_lines.append(" | ".join(cell_values))

        return TableContent(
            table_number=table_num,
            context_heading=context,
            text_content="\n".join(text_lines),
            row_count=len(text_lines),
        )

    def _is_heading(self, text: str) -> bool:
        """Check if text looks like a heading (e.g., '1.1 Purpose')"""
        if not text:
            return False
        return bool(re.match(r"^\d+(\.\d+)*\s+[A-Z]", text.strip()))

    def _extract_project_id(self, file_path: Path) -> str:
        """Extract project ID from folder name or file name"""
        # Try folder name first (e.g., PRJ-001)
        folder_name = file_path.parent.name
        match = re.match(r"(PRJ-\d+)", folder_name)
        if match:
            return match.group(1)

        # Try file name
        file_stem = file_path.stem
        match = re.match(r"(PRJ-\d+)", file_stem)
        if match:
            return match.group(1)

        # Fall back to folder name
        return folder_name
