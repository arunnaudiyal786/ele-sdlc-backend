"""
Project Indexer Service

Creates and manages a lightweight index of project metadata for fast hybrid search.
Only indexes project_name + summary for embedding; full documents loaded on-demand.
"""

import logging
import re
import threading
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from pydantic import BaseModel, Field

from app.rag.vector_store import ChromaVectorStore
from app.rag.embeddings import OllamaEmbeddingService
from app.components.base.config import get_settings

logger = logging.getLogger(__name__)


class ProjectMetadata(BaseModel):
    """Lightweight metadata for project index"""

    project_id: str = Field(..., description="Project ID, e.g., PRJ-10051")
    project_name: str = Field(..., description="Human-readable project name")
    summary: str = Field(..., description="Epic description from TDD 1.1 Purpose")
    folder_path: str = Field(..., description="Absolute path to project folder")
    tdd_path: str = Field(..., description="Path to TDD.docx")
    estimation_path: str = Field(..., description="Path to estimation.xlsx")
    jira_stories_path: str = Field(..., description="Path to jira_stories.xlsx")
    indexed_at: datetime = Field(default_factory=datetime.now)

    @property
    def document_text(self) -> str:
        """Text used for embedding in ChromaDB"""
        return f"{self.project_name} {self.summary}"


class ProjectIndexer:
    """
    Manages lightweight project metadata index in ChromaDB.

    Scans project folders, extracts metadata from TDD documents,
    and maintains a searchable index for hybrid search.
    """

    _instance: Optional["ProjectIndexer"] = None
    _lock = threading.Lock()

    def __init__(
        self,
        vector_store: Optional[ChromaVectorStore] = None,
        embedding_service: Optional[OllamaEmbeddingService] = None,
    ):
        settings = get_settings()
        self.vector_store = vector_store or ChromaVectorStore(
            persist_dir=settings.chroma_persist_dir
        )
        self.embedding_service = embedding_service or OllamaEmbeddingService()
        self.collection_name = "project_index"

    @classmethod
    def get_instance(cls) -> "ProjectIndexer":
        """Thread-safe singleton instance"""
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = cls()
        return cls._instance

    async def scan_projects(self, base_path: Path) -> List[ProjectMetadata]:
        """
        Scan data/raw/projects/ folder for project directories

        Args:
            base_path: Path to projects root (e.g., data/raw/projects)

        Returns:
            List of ProjectMetadata for all valid projects
        """
        logger.info(f"Scanning projects in {base_path}")

        if not base_path.exists():
            logger.error(f"Project base path does not exist: {base_path}")
            return []

        projects = []

        # Iterate through subdirectories
        for project_folder in base_path.iterdir():
            # Skip files (like epic.csv)
            if not project_folder.is_dir():
                continue

            # Skip hidden directories
            if project_folder.name.startswith("."):
                continue

            try:
                metadata = await self.extract_metadata(project_folder)
                projects.append(metadata)
                logger.info(f"Found project: {metadata.project_id}")
            except Exception as e:
                logger.warning(
                    f"Failed to extract metadata from {project_folder.name}: {e}"
                )
                continue

        logger.info(f"Scanned {len(projects)} projects")
        return projects

    async def extract_metadata(self, project_folder: Path) -> ProjectMetadata:
        """
        Extract metadata from TDD.docx file

        Strategy:
        - project_id: Parse from folder name (PRJ-XXXXX pattern)
        - project_name: Extract from TDD References section or derive from folder
        - summary: Extract from TDD section 1.1 Purpose (paragraph after heading)

        Args:
            project_folder: Path to project directory

        Returns:
            ProjectMetadata with extracted information

        Raises:
            FileNotFoundError: If TDD file not found
            ValueError: If metadata extraction fails
        """
        tdd_path = project_folder / "tdd.docx"
        estimation_path = project_folder / "estimation.xlsx"
        jira_stories_path = project_folder / "jira_stories.xlsx"

        # Verify required files exist
        if not tdd_path.exists():
            raise FileNotFoundError(f"TDD file not found: {tdd_path}")

        # Extract project_id from folder name
        project_id = self._extract_project_id(project_folder.name)

        # Parse TDD document
        doc = Document(str(tdd_path))

        # Extract project_name and summary
        project_name = self._extract_project_name(doc, project_folder.name)
        summary = self._extract_purpose(doc)

        return ProjectMetadata(
            project_id=project_id,
            project_name=project_name,
            summary=summary,
            folder_path=str(project_folder.absolute()),
            tdd_path=str(tdd_path.absolute()),
            estimation_path=str(estimation_path.absolute()),
            jira_stories_path=str(jira_stories_path.absolute()),
            indexed_at=datetime.now(),
        )

    def _extract_project_id(self, folder_name: str) -> str:
        """
        Extract project ID from folder name (PRJ-XXXXX pattern)

        Examples:
            PRJ-10051-inventory-sync-automation → PRJ-10051
            PRJ-10052-order-fulfillment-optimization → PRJ-10052

        Args:
            folder_name: Name of project folder

        Returns:
            Project ID in format PRJ-XXXXX
        """
        # Match PRJ-XXXXX pattern at the start of folder name
        match = re.match(r"^(PRJ-\d+)", folder_name)
        if match:
            return match.group(1)

        # Fallback: try to get first two parts (PREFIX-NUMBER)
        parts = folder_name.split("-")
        if len(parts) >= 2 and parts[1].isdigit():
            return f"{parts[0]}-{parts[1]}"

        # Last resort: use entire folder name
        logger.warning(f"Could not extract project ID from '{folder_name}', using entire name")
        return folder_name

    def _extract_project_name(self, doc: Document, folder_name: str) -> str:
        """
        Extract project name from TDD References section

        Strategy:
        1. Look for "PRJ-XXXXX Project Charter" pattern in References section
        2. Extract project name from that line
        3. Fallback: Derive from folder name (replace hyphens with spaces, title case)

        Args:
            doc: TDD Word document
            folder_name: Folder name as fallback

        Returns:
            Human-readable project name
        """
        # Search for References section (section 1.2)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Look for "PRJ-XXXXX Project Charter" or similar
            if "PRJ-" in text and ("Project Charter" in text or "Initiative" in text):
                # Extract project name after PRJ-XXXXX
                # Example: "PRJ-10051 Project Charter" → "Project Charter"
                # Or: "INV-2026 - Inventory Sync Real-Time Initiative" → "Inventory Sync Real-Time Initiative"
                match = re.search(r"PRJ-\d+\s*[-:]?\s*(.+?)(?:Project Charter|Initiative)?$", text)
                if match:
                    name = match.group(1).strip().strip("-").strip()
                    if name:
                        return name

        # Fallback: Derive from folder name
        # PRJ-10051-inventory-sync-automation → Inventory Sync Automation
        project_id_match = re.match(r"PRJ-\d+-(.+)", folder_name)
        if project_id_match:
            name_part = project_id_match.group(1)
            # Replace hyphens with spaces and title case
            project_name = name_part.replace("-", " ").title()
            logger.info(
                f"Using derived project name from folder: {project_name}"
            )
            return project_name

        # Last resort: use folder name
        return folder_name.replace("-", " ").title()

    def _extract_purpose(self, doc: Document) -> str:
        """
        Extract epic description from 1.1 Purpose section

        Strategy:
        1. Find heading containing "1.1" and "Purpose"
        2. Return the next non-empty paragraph
        3. Fallback: Return first substantial paragraph (>50 chars)

        Args:
            doc: TDD Word document

        Returns:
            Epic description text
        """
        # Strategy 1: Find "1.1 Purpose" heading and get next paragraph
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Check if this is the Purpose heading
            if "1.1" in text and "Purpose" in text.lower():
                # Get next non-empty paragraph
                for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j].text.strip()
                    if len(next_para) > 20:  # Meaningful content
                        return next_para

        # Strategy 2: Find first paragraph after "INTRODUCTION" or "Purpose"
        found_intro = False
        for para in doc.paragraphs:
            text = para.text.strip()

            if "INTRODUCTION" in text.upper() or "PURPOSE" in text.upper():
                found_intro = True
                continue

            if found_intro and len(text) > 50:
                return text

        # Fallback: Return first substantial paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 50 and not text.isupper():  # Not a heading
                logger.warning("Using fallback purpose extraction")
                return text

        logger.error("Could not extract purpose from TDD")
        return "No purpose section found"

    async def build_index(self, base_path: Path) -> int:
        """
        Build complete index from scratch

        Scans all projects, extracts metadata, embeds documents,
        and stores in ChromaDB collection.

        Args:
            base_path: Path to projects root

        Returns:
            Number of projects indexed
        """
        logger.info(f"Building project index from {base_path}")

        # Scan all projects
        projects = await self.scan_projects(base_path)

        if not projects:
            logger.warning("No projects found to index")
            return 0

        # Create or recreate collection
        try:
            # Delete existing collection if it exists
            await self.vector_store.delete_collection(self.collection_name)
            logger.info(f"Deleted existing collection: {self.collection_name}")
        except Exception:
            pass  # Collection doesn't exist yet

        # Collection will be auto-created on first add via get_or_create_collection
        logger.info(f"Preparing collection: {self.collection_name}")

        # Index each project
        indexed_count = 0
        for project in projects:
            try:
                await self._index_project(project)
                indexed_count += 1
                logger.info(f"Indexed: {project.project_id}")
            except Exception as e:
                logger.error(f"Failed to index {project.project_id}: {e}")
                continue

        logger.info(f"Successfully indexed {indexed_count}/{len(projects)} projects")
        return indexed_count

    async def _index_project(self, metadata: ProjectMetadata) -> None:
        """
        Index a single project in ChromaDB

        Args:
            metadata: Project metadata to index
        """
        # Generate embedding for document text
        embedding = await self.embedding_service.embed(metadata.document_text)

        # Prepare metadata dict for ChromaDB
        chroma_metadata = {
            "project_id": metadata.project_id,
            "project_name": metadata.project_name,
            "summary": metadata.summary,
            "folder_path": metadata.folder_path,
            "tdd_path": metadata.tdd_path,
            "estimation_path": metadata.estimation_path,
            "jira_stories_path": metadata.jira_stories_path,
            "indexed_at": metadata.indexed_at.isoformat(),
        }

        # Prepare document in format expected by add_documents
        document = {
            "id": metadata.project_id,
            "text": metadata.document_text,
            "metadata": chroma_metadata,
        }

        # Add to ChromaDB collection
        await self.vector_store.add_documents(
            collection_name=self.collection_name,
            documents=[document],
            embeddings=[embedding],
        )

    async def add_project(self, project_folder: Path) -> str:
        """
        Add single project to index

        Args:
            project_folder: Path to project directory

        Returns:
            Project ID of added project

        Raises:
            ValueError: If project extraction or indexing fails
        """
        logger.info(f"Adding project: {project_folder}")

        # Extract metadata
        metadata = await self.extract_metadata(project_folder)

        # Index project
        await self._index_project(metadata)

        logger.info(f"Added project: {metadata.project_id}")
        return metadata.project_id

    async def refresh_index(self, base_path: Path) -> int:
        """
        Refresh existing index (update modified projects)

        Currently rebuilds entire index. Future optimization:
        - Track modification times
        - Only re-index changed projects

        Args:
            base_path: Path to projects root

        Returns:
            Number of projects re-indexed
        """
        logger.info("Refreshing project index (full rebuild)")
        return await self.build_index(base_path)

    async def remove_project(self, project_id: str) -> bool:
        """
        Remove project from index

        Args:
            project_id: Project ID to remove

        Returns:
            True if successfully removed, False otherwise
        """
        try:
            collection = self.vector_store.get_or_create_collection(self.collection_name)
            collection.delete(ids=[project_id])
            logger.info(f"Removed project: {project_id}")
            return True
        except Exception as e:
            logger.error(f"Failed to remove project {project_id}: {e}")
            return False
